from __future__ import annotations

import io
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from loguru import logger

# Lazy imports agar tidak crash bila salah satu lib belum di-install
try:
    import fitz  # PyMuPDF
    _PYMUPDF_AVAILABLE = True
except ImportError:
    _PYMUPDF_AVAILABLE = False
    logger.warning("PyMuPDF (fitz) tidak terinstall. Pure PDF tidak akan bisa dibaca.")

try:
    from PIL import Image
    _PIL_AVAILABLE = True
except ImportError:
    _PIL_AVAILABLE = False

try:
    from paddleocr import PaddleOCR
    _PADDLEOCR_AVAILABLE = True
except ImportError:
    _PADDLEOCR_AVAILABLE = False
    logger.warning("PaddleOCR tidak terinstall. Scanned PDF tidak akan bisa dibaca.")

try:
    from docx import Document as DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False
    logger.warning("python-docx tidak terinstall. DOCX tidak akan bisa dibaca.")


# ─────────────────────────────────────────────────────────────
# Data model
# ─────────────────────────────────────────────────────────────
@dataclass
class PageText:
    page_number: int
    text: str
    source: str          # "pure" | "ocr" | "docx"
    confidence: float    # 1.0 untuk pure/docx, rata-rata confidence OCR untuk scanned


# ─────────────────────────────────────────────────────────────
# OCR singleton (inisialisasi mahal, reuse satu instance)
# ─────────────────────────────────────────────────────────────
_ocr_instance: Optional["PaddleOCR"] = None

def _get_ocr(lang: str = "id", use_gpu: bool = False) -> "PaddleOCR":
    global _ocr_instance
    if _ocr_instance is None:
        if not _PADDLEOCR_AVAILABLE:
            raise RuntimeError("PaddleOCR tidak terinstall. Jalankan: pip install paddleocr")
        logger.info("Inisialisasi PaddleOCR...")
        _ocr_instance = PaddleOCR(
            use_angle_cls=True,
            lang=lang,
            use_gpu=use_gpu,
            show_log=False,
        )
    return _ocr_instance


# ─────────────────────────────────────────────────────────────
# PDF type detection
# ─────────────────────────────────────────────────────────────
def is_scanned_pdf(pdf_path: str | Path, char_threshold: int = 50) -> bool:
    """
    Deteksi apakah PDF merupakan hasil scan (tidak punya teks selectable).
    
    Strategi: buka dokumen, hitung total karakter teks dari semua halaman.
    Jika < char_threshold → dianggap scanned.
    """
    if not _PYMUPDF_AVAILABLE:
        logger.warning("PyMuPDF tidak tersedia, asumsikan scanned PDF.")
        return True

    doc = fitz.open(str(pdf_path))
    total_chars = sum(len(page.get_text().strip()) for page in doc)
    doc.close()

    result = total_chars < char_threshold
    logger.debug(
        f"PDF '{Path(pdf_path).name}': total_chars={total_chars}, "
        f"threshold={char_threshold}, is_scanned={result}"
    )
    return result


# ─────────────────────────────────────────────────────────────
# Pure PDF extractor (PyMuPDF) — dengan crop header/footer
# ─────────────────────────────────────────────────────────────
def extract_pure_pdf(
    pdf_path: str | Path,
    skip_header_footer: bool = True,
    header_ratio: float = 0.12,
    footer_ratio: float = 0.10,
) -> List[PageText]:
    """
    Ekstrak teks dari PDF pure menggunakan PyMuPDF.

    Jika skip_header_footer=True, area atas (header_ratio) dan bawah
    (footer_ratio) dari setiap halaman dipotong agar kop surat dan
    footer institusi tidak ikut terbaca.
    """
    if not _PYMUPDF_AVAILABLE:
        raise RuntimeError("PyMuPDF tidak terinstall. Jalankan: pip install pymupdf")

    results: List[PageText] = []
    doc = fitz.open(str(pdf_path))

    for page_num, page in enumerate(doc, start=1):
        if skip_header_footer and (header_ratio > 0 or footer_ratio > 0):
            # Hitung clip rectangle — potong area header dan footer
            rect = page.rect                     # fitz.Rect(x0, y0, x1, y1)
            page_height = rect.height
            y_top    = rect.y0 + page_height * header_ratio
            y_bottom = rect.y1 - page_height * footer_ratio
            clip = fitz.Rect(rect.x0, y_top, rect.x1, y_bottom)

            raw_text = page.get_text("text", sort=True, clip=clip)
            logger.debug(
                f"  [Pure] Hal. {page_num}: clip y={y_top:.0f}-{y_bottom:.0f} "
                f"(header={header_ratio*100:.0f}%, footer={footer_ratio*100:.0f}%)"
            )
        else:
            raw_text = page.get_text("text", sort=True)

        cleaned = _clean_text(raw_text)
        results.append(PageText(
            page_number=page_num,
            text=cleaned,
            source="pure",
            confidence=1.0,
        ))
        logger.debug(f"  [Pure] Hal. {page_num}: {len(cleaned)} karakter")

    doc.close()
    return results


# ─────────────────────────────────────────────────────────────
# Scanned PDF extractor (PyMuPDF render → PaddleOCR)
# — dengan crop header/footer
# ─────────────────────────────────────────────────────────────
def extract_scanned_pdf(
    pdf_path: str | Path,
    dpi: int = 200,
    lang: str = "id",
    use_gpu: bool = False,
    max_img_size: int = 4096,
    skip_header_footer: bool = True,
    header_ratio: float = 0.12,
    footer_ratio: float = 0.10,
) -> List[PageText]:
    """
    Render setiap halaman PDF ke gambar, lalu jalankan PaddleOCR.

    Jika skip_header_footer=True, gambar di-crop untuk menghilangkan
    area header dan footer sebelum OCR dijalankan.
    """
    if not _PYMUPDF_AVAILABLE:
        raise RuntimeError("PyMuPDF tidak terinstall.")
    if not _PIL_AVAILABLE:
        raise RuntimeError("Pillow tidak terinstall. Jalankan: pip install Pillow")

    ocr = _get_ocr(lang=lang, use_gpu=use_gpu)
    results: List[PageText] = []

    doc  = fitz.open(str(pdf_path))
    zoom = dpi / 72.0             # 72 DPI adalah default PDF
    mat  = fitz.Matrix(zoom, zoom)

    for page_num, page in enumerate(doc, start=1):
        logger.debug(f"  [OCR] Render hal. {page_num}...")

        # Render ke pixmap
        pix = page.get_pixmap(matrix=mat, alpha=False)
        img_bytes = pix.tobytes("png")

        # Load ke PIL, resize jika terlalu besar
        img = Image.open(io.BytesIO(img_bytes)).convert("RGB")

        # Crop header/footer dari gambar
        if skip_header_footer and (header_ratio > 0 or footer_ratio > 0):
            w, h = img.size
            y_top    = int(h * header_ratio)
            y_bottom = int(h * (1 - footer_ratio))
            img = img.crop((0, y_top, w, y_bottom))
            logger.debug(
                f"    Crop header/footer: y={y_top}-{y_bottom} "
                f"(header={header_ratio*100:.0f}%, footer={footer_ratio*100:.0f}%)"
            )

        if max(img.size) > max_img_size:
            ratio = max_img_size / max(img.size)
            new_size = (int(img.width * ratio), int(img.height * ratio))
            img = img.resize(new_size, Image.LANCZOS)
            logger.debug(f"    Resize gambar ke {new_size}")

        # Konversi kembali ke bytes untuk PaddleOCR
        buf = io.BytesIO()
        img.save(buf, format="PNG")
        img_array = _pil_to_numpy(img)

        # Jalankan OCR
        ocr_result = ocr.ocr(img_array, cls=True)

        page_text, confidences = _parse_paddle_result(ocr_result)
        avg_conf = sum(confidences) / len(confidences) if confidences else 0.0

        results.append(PageText(
            page_number=page_num,
            text=_clean_text(page_text),
            source="ocr",
            confidence=avg_conf,
        ))
        logger.debug(
            f"  [OCR] Hal. {page_num}: {len(page_text)} karakter, "
            f"conf={avg_conf:.3f}"
        )

    doc.close()
    return results


def _parse_paddle_result(ocr_result) -> tuple[str, list[float]]:
    """Parse output PaddleOCR menjadi string dan list confidence."""
    lines, confidences = [], []

    if ocr_result is None:
        return "", []

    for page_data in ocr_result:
        if page_data is None:
            continue
        for line in page_data:
            # line = [ [[x1,y1],...], (text, confidence) ]
            if len(line) >= 2:
                text_info = line[1]
                text  = text_info[0] if isinstance(text_info, (list, tuple)) else str(text_info)
                conf  = float(text_info[1]) if isinstance(text_info, (list, tuple)) and len(text_info) > 1 else 1.0
                lines.append(text)
                confidences.append(conf)

    return "\n".join(lines), confidences


def _pil_to_numpy(img: "Image.Image"):
    """Konversi PIL Image ke numpy array (diperlukan PaddleOCR)."""
    import numpy as np
    return np.array(img)


def _clean_text(text: str) -> str:
    """Normalisasi whitespace pada teks hasil ekstraksi."""
    import re
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    text = re.sub(r"\n{3,}", "\n\n", text)    # max 2 newline berturut-turut
    text = re.sub(r"[ \t]+", " ", text)       # spasi ganda
    return text.strip()


# ─────────────────────────────────────────────────────────────
# DOCX extractor (python-docx) — otomatis skip header/footer
# ─────────────────────────────────────────────────────────────
def extract_text_from_docx(docx_path: str | Path) -> List[PageText]:
    """
    Ekstrak teks dari file DOCX menggunakan python-docx.

    DOCX menyimpan header dan footer secara terpisah dari body document,
    sehingga secara default hanya body paragraphs yang dibaca (header
    dan footer otomatis ter-skip).

    Tabel dalam DOCX juga diekstrak dan digabungkan ke teks.

    Returns:
        List[PageText] — karena DOCX tidak punya konsep "halaman" secara
        native, seluruh isi dikembalikan sebagai satu PageText (page 1).
    """
    if not _DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx tidak terinstall. Jalankan: pip install python-docx"
        )

    docx_path = Path(docx_path)
    logger.info(f"Membaca DOCX: {docx_path.name}")

    doc = DocxDocument(str(docx_path))

    # ── Ekstrak paragraf body (otomatis tanpa header/footer) ──
    paragraphs: list[str] = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # ── Ekstrak tabel ────────────────────────────────────────
    table_texts: list[str] = []
    for table_idx, table in enumerate(doc.tables):
        rows_text: list[str] = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            # Hapus duplikat sel yang terjadi karena merged cells
            unique_cells: list[str] = []
            for c in cells:
                if c and (not unique_cells or c != unique_cells[-1]):
                    unique_cells.append(c)
            if unique_cells:
                rows_text.append(" | ".join(unique_cells))
        if rows_text:
            table_texts.append(
                f"[Tabel {table_idx + 1}]\n" + "\n".join(rows_text)
            )

    # ── Gabungkan paragraf dan tabel ─────────────────────────
    body_text = "\n".join(paragraphs)
    if table_texts:
        body_text += "\n\n[KONTEKS TABEL]\n" + "\n\n".join(table_texts)

    cleaned = _clean_text(body_text)

    result = [
        PageText(
            page_number=1,
            text=cleaned,
            source="docx",
            confidence=1.0,
        )
    ]

    logger.info(f"  ✓ DOCX berhasil dibaca: {len(cleaned)} karakter, "
                f"{len(paragraphs)} paragraf, {len(doc.tables)} tabel")

    return result


# ─────────────────────────────────────────────────────────────
# Main entry point — PDF
# ─────────────────────────────────────────────────────────────
def extract_text_from_pdf(
    pdf_path: str | Path,
    char_threshold: int = 50,
    dpi: int = 200,
    lang: str = "id",
    use_gpu: bool = False,
    skip_header_footer: bool = True,
    header_ratio: float = 0.12,
    footer_ratio: float = 0.10,
) -> List[PageText]:
    pdf_path = Path(pdf_path)
    if not pdf_path.exists():
        raise FileNotFoundError(f"File PDF tidak ditemukan: {pdf_path}")

    logger.info(f"Memproses PDF: {pdf_path.name}")

    if is_scanned_pdf(pdf_path, char_threshold=char_threshold):
        logger.info("→ Jenis PDF: SCANNED → menggunakan PaddleOCR")
        return extract_scanned_pdf(
            pdf_path, dpi=dpi, lang=lang, use_gpu=use_gpu,
            skip_header_footer=skip_header_footer,
            header_ratio=header_ratio,
            footer_ratio=footer_ratio,
        )
    else:
        logger.info("→ Jenis PDF: PURE → menggunakan PyMuPDF")
        return extract_pure_pdf(
            pdf_path,
            skip_header_footer=skip_header_footer,
            header_ratio=header_ratio,
            footer_ratio=footer_ratio,
        )


# ─────────────────────────────────────────────────────────────
# Dispatcher: PDF atau DOCX
# ─────────────────────────────────────────────────────────────
def extract_text_from_document(
    file_path: str | Path,
    char_threshold: int = 50,
    dpi: int = 200,
    lang: str = "id",
    use_gpu: bool = False,
    skip_header_footer: bool = True,
    header_ratio: float = 0.12,
    footer_ratio: float = 0.10,
) -> List[PageText]:
    """
    Dispatcher utama: deteksi tipe file (PDF/DOCX) dan ekstrak teks.

    - PDF  → extract_text_from_pdf() dengan header/footer cropping
    - DOCX → extract_text_from_docx() (header/footer otomatis ter-skip)

    Args:
        file_path: Path ke file dokumen (.pdf atau .docx)
        char_threshold: Threshold karakter untuk menentukan PDF pure vs scanned
        dpi: DPI render untuk scanned PDF
        lang: Bahasa PaddleOCR
        use_gpu: Gunakan GPU untuk PaddleOCR
        skip_header_footer: Potong header/footer (hanya berlaku untuk PDF)
        header_ratio: Rasio area atas yang dipotong (0.0-1.0)
        footer_ratio: Rasio area bawah yang dipotong (0.0-1.0)

    Returns:
        List[PageText] — teks per halaman (atau satu elemen untuk DOCX)
    """
    file_path = Path(file_path)
    if not file_path.exists():
        raise FileNotFoundError(f"File tidak ditemukan: {file_path}")

    ext = file_path.suffix.lower()

    if ext == ".pdf":
        return extract_text_from_pdf(
            file_path,
            char_threshold=char_threshold,
            dpi=dpi,
            lang=lang,
            use_gpu=use_gpu,
            skip_header_footer=skip_header_footer,
            header_ratio=header_ratio,
            footer_ratio=footer_ratio,
        )
    elif ext == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(
            f"Tipe file tidak didukung: '{ext}'. "
            f"Gunakan .pdf atau .docx"
        )


def pages_to_full_text(pages: List[PageText]) -> str:
    """Gabungkan semua halaman menjadi satu string."""
    return "\n\n".join(
        f"[Halaman {p.page_number}]\n{p.text}"
        for p in pages
        if p.text.strip()
    )


def extract_text_and_tables(
    pdf_path: str | Path,
    char_threshold: int  = 50,
    dpi: int             = 200,
    lang: str            = "id",
    use_gpu: bool        = False,
    table_style: str     = "kv",
    merge_tables: bool   = True,
    skip_header_footer: bool = True,
    header_ratio: float  = 0.12,
    footer_ratio: float  = 0.10,
    
) -> dict:
    from src.table_extractor import extract_tables_from_pdf, tables_to_ner_context

    pdf_path   = Path(pdf_path)
    scanned    = is_scanned_pdf(pdf_path, char_threshold=char_threshold)

    # Ekstrak teks biasa
    pages      = extract_text_from_pdf(
        pdf_path, char_threshold=char_threshold,
        dpi=dpi, lang=lang, use_gpu=use_gpu,
        skip_header_footer=skip_header_footer,
        header_ratio=header_ratio,
        footer_ratio=footer_ratio,
    )
    full_text  = pages_to_full_text(pages)

    # Ekstrak tabel
    tables, table_text = extract_tables_from_pdf(
        pdf_path,
        is_scanned    = scanned,
        dpi           = dpi,
        lang          = lang,
        use_gpu       = use_gpu,
        table_text_style = table_style,
        fallback_text = full_text,   # fallback heuristik jika metode utama gagal
    )

    # Gabungkan teks biasa + teks tabel untuk konteks NER
    if merge_tables and table_text.strip():
        combined = full_text + "\n\n[KONTEKS TABEL]\n" + table_text
    else:
        combined = full_text

    return {
        "pages":      pages,
        "tables":     tables,
        "full_text":  full_text,
        "table_text": table_text,
        "combined":   combined,
        "is_scanned": scanned,
        "n_tables":   len(tables),
    }